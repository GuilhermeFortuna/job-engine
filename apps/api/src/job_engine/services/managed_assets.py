"""Managed local storage service for multi-profile assets."""

from __future__ import annotations

import hashlib
import io
import os
import re
import zipfile
from collections.abc import AsyncIterator
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import docx
from pypdf import PdfReader

from job_engine.domain.applicant import AvatarCrop, ManagedAsset, ManagedAssetType

MAX_RESUME_SIZE_BYTES = 20 * 1024 * 1024  # 20 MiB
MAX_AVATAR_SIZE_BYTES = 10 * 1024 * 1024  # 10 MiB
CHUNK_SIZE = 64 * 1024  # 64 KiB

PNG_MAGIC = b"\x89PNG\r\n\x1a\n"
JPEG_MAGIC = b"\xff\xd8\xff"
PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"


class InvalidAssetError(ValueError):
    """Raised when an uploaded asset violates format, type, or integrity rules."""


class InvalidAssetTypeError(InvalidAssetError):
    """Raised when an asset type or magic header does not match expected format."""


class OversizeAssetError(InvalidAssetError):
    """Raised when an uploaded asset exceeds maximum allowed size."""


class AssetNotFoundError(LookupError):
    """Raised when an asset or its physical file cannot be found."""


def sanitize_filename(name: str, fallback_ext: str = "") -> str:
    """Produce a safe, single-level filename free of path traversal elements."""
    raw_name = Path(name).name
    # Strip null bytes and control chars
    clean = re.sub(r"[\x00-\x1f\x7f/\\]", "", raw_name).strip()
    # Replace dangerous or unwanted chars
    clean = re.sub(r"[^a-zA-Z0-9_.-]", "_", clean)
    # Avoid leading dots or empty names
    clean = clean.lstrip(".")
    if not clean:
        ext = fallback_ext.lstrip(".") if fallback_ext else "bin"
        clean = f"asset.{ext}"
    return clean


def detect_file_type(
    header: bytes, declared_content_type: str | None = None
) -> tuple[str, str]:
    """Detect format from magic bytes and validate against declared content type."""
    if header.startswith(PDF_MAGIC):
        return "application/pdf", ".pdf"
    if header.startswith(PNG_MAGIC):
        return "image/png", ".png"
    if header.startswith(JPEG_MAGIC):
        return "image/jpeg", ".jpg"
    if len(header) >= 12 and header.startswith(b"RIFF") and header[8:12] == b"WEBP":
        return "image/webp", ".webp"
    if header.startswith(ZIP_MAGIC):
        # Check if it's a docx
        try:
            with zipfile.ZipFile(io.BytesIO(header)) as zf:
                namelist = zf.namelist()
                if (
                    any(n.startswith("word/") for n in namelist)
                    or "[Content_Types].xml" in namelist
                ):
                    return (
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ".docx",
                    )
        except Exception:
            pass
        return (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx",
        )

    raise InvalidAssetTypeError(
        "Unrecognized or unsupported file signature "
        f"(declared: {declared_content_type})"
    )


def validate_resume_or_document_type(
    header: bytes, content_type: str | None, filename: str
) -> tuple[str, str]:
    """Ensure asset is PDF or DOCX."""
    detected_mime, detected_ext = detect_file_type(header, content_type)
    if detected_mime not in {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        raise InvalidAssetTypeError(f"Expected PDF or DOCX, detected {detected_mime}")
    return detected_mime, detected_ext


def validate_avatar_type(
    header: bytes, content_type: str | None, filename: str
) -> tuple[str, str]:
    """Ensure asset is PNG, JPEG, or WebP."""
    detected_mime, detected_ext = detect_file_type(header, content_type)
    if detected_mime not in {"image/png", "image/jpeg", "image/webp"}:
        raise InvalidAssetTypeError(
            f"Expected PNG, JPEG, or WebP avatar image, detected {detected_mime}"
        )
    return detected_mime, detected_ext


def extract_text_from_file(file_path: Path, content_type: str) -> str | None:
    """Extract plain text from PDF or DOCX file."""
    try:
        if content_type == "application/pdf":
            reader = PdfReader(str(file_path))
            pages_text: list[str] = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted.strip())
            return "\n\n".join(pages_text).strip() or None
        elif content_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/docx",
        }:
            doc = docx.Document(str(file_path))
            parts: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for t in doc.tables:
                for row in t.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts).strip() or None
    except Exception:
        return None
    return None


class ManagedAssetService:
    def __init__(self, data_root: Path) -> None:
        self._data_root = Path(data_root).expanduser().resolve()

    @property
    def data_root(self) -> Path:
        return self._data_root

    def ensure_root(self) -> None:
        self._data_root.mkdir(parents=True, exist_ok=True)

    def resolve_physical_path(self, relative_path: str) -> Path:
        """Resolve a stored relative path safely, ensuring no traversal escapes."""
        # Normalize relative path
        norm_rel = Path(relative_path)
        if norm_rel.is_absolute():
            raise InvalidAssetError(
                f"Absolute path not permitted in managed asset: {relative_path}"
            )

        target = (self._data_root / norm_rel).resolve()

        # Ensure target is strictly beneath data_root
        try:
            target.relative_to(self._data_root)
        except ValueError as exc:
            raise InvalidAssetError(
                f"Path traversal detected: {relative_path}"
            ) from exc

        return target

    async def store_asset_stream(
        self,
        profile_id: UUID,
        asset_type: ManagedAssetType,
        filename: str,
        content_stream: AsyncIterator[bytes],
        declared_content_type: str | None = None,
        crop_coordinates: AvatarCrop | dict[str, Any] | None = None,
    ) -> ManagedAsset:
        """Stream bytes to a sibling temporary file and return domain model."""
        self.ensure_root()

        max_size = (
            MAX_AVATAR_SIZE_BYTES
            if asset_type == ManagedAssetType.AVATAR
            else MAX_RESUME_SIZE_BYTES
        )
        asset_id = uuid4()

        safe_name = sanitize_filename(filename)
        asset_dir = (
            self._data_root / "profiles" / str(profile_id) / "assets" / str(asset_id)
        )
        asset_dir.mkdir(parents=True, exist_ok=True)

        temp_file_path = asset_dir / f".tmp_{uuid4().hex}"
        hasher = hashlib.sha256()
        total_bytes = 0
        header_bytes = bytearray()
        validated_type = False
        final_mime = declared_content_type or "application/octet-stream"

        try:
            with open(temp_file_path, "wb") as f:  # noqa: ASYNC230
                async for chunk in content_stream:
                    if not chunk:
                        continue
                    total_bytes += len(chunk)
                    if total_bytes > max_size:
                        raise OversizeAssetError(
                            f"Asset exceeds maximum allowed size of {max_size} bytes "
                            f"(got {total_bytes})"
                        )
                    hasher.update(chunk)

                    if len(header_bytes) < 4096:
                        header_bytes.extend(chunk[: 4096 - len(header_bytes)])

                    if not validated_type and len(header_bytes) >= 16:
                        if asset_type == ManagedAssetType.AVATAR:
                            final_mime, _ = validate_avatar_type(
                                bytes(header_bytes), declared_content_type, safe_name
                            )
                        else:
                            final_mime, _ = validate_resume_or_document_type(
                                bytes(header_bytes), declared_content_type, safe_name
                            )
                        validated_type = True

                    f.write(chunk)

                if total_bytes == 0:
                    raise InvalidAssetError("Uploaded file is empty (0 bytes)")

                if not validated_type:
                    if asset_type == ManagedAssetType.AVATAR:
                        final_mime, _ = validate_avatar_type(
                            bytes(header_bytes), declared_content_type, safe_name
                        )
                    else:
                        final_mime, _ = validate_resume_or_document_type(
                            bytes(header_bytes), declared_content_type, safe_name
                        )

                f.flush()
                os.fsync(f.fileno())

            final_file_path = asset_dir / safe_name
            os.replace(temp_file_path, final_file_path)

            sha256_hex = hasher.hexdigest()
            rel_path = f"profiles/{profile_id}/assets/{asset_id}/{safe_name}"

            extracted_text = None
            if asset_type in {ManagedAssetType.RESUME, ManagedAssetType.DOCUMENT}:
                extracted_text = extract_text_from_file(final_file_path, final_mime)

            crop_dict = None
            if crop_coordinates is not None:
                if isinstance(crop_coordinates, AvatarCrop):
                    crop_dict = crop_coordinates.model_dump()
                elif isinstance(crop_coordinates, dict):
                    crop_dict = crop_coordinates

            now = datetime.now(UTC)
            return ManagedAsset(
                id=asset_id,
                profile_id=profile_id,
                asset_type=asset_type,
                file_name=safe_name,
                content_type=final_mime,
                byte_size=total_bytes,
                sha256=sha256_hex,
                relative_path=rel_path,
                crop_coordinates=crop_dict,
                extracted_text=extracted_text,
                created_at=now,
                updated_at=now,
            )
        except Exception:
            # Clean up on failure
            if temp_file_path.exists():
                try:
                    temp_file_path.unlink()
                except OSError:
                    pass
            try:
                # If directory is empty, remove it
                if asset_dir.exists() and not any(asset_dir.iterdir()):
                    asset_dir.rmdir()
            except OSError:
                pass
            raise

    def get_asset_file(self, relative_path: str) -> tuple[Path, int]:
        """Verify and return physical path and file size."""
        target = self.resolve_physical_path(relative_path)
        if not target.is_file():
            raise AssetNotFoundError(f"Asset file not found on disk: {relative_path}")
        return target, target.stat().st_size

    def delete_asset_file(self, relative_path: str) -> None:
        """Unlink file and parent asset directory if empty."""
        try:
            target = self.resolve_physical_path(relative_path)
            if target.exists() and target.is_file():
                target.unlink()
            parent = target.parent
            if parent.exists() and not any(parent.iterdir()):
                parent.rmdir()
        except Exception:
            pass
